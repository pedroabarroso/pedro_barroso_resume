from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

path = r"C:\Users\pedro\repos\resume\Pedro_Barroso_Data_Analyst.docx"
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, size=12, bold=True, color=(0x6E, 0x53, 0x38))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "8B6B4A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_para(text, size=11, bold=False, italic=False, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(2)
set_run_font(name.add_run("PEDRO BARROSO"), size=20, bold=True)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(2)
set_run_font(
    contact.add_run(
        "Liwan, Dubai, UAE  |  WhatsApp +971 52 279 625 (wa.me/97152279625)  |  "
        "pedrobarroso9mmm@gmail.com  |  linkedin.com/in/pedro-a-barroso"
    ),
    size=10,
)

role = doc.add_paragraph()
role.alignment = WD_ALIGN_PARAGRAPH.CENTER
role.paragraph_format.space_after = Pt(8)
set_run_font(role.add_run("Data Analyst"), size=11, bold=True, color=(0x8B, 0x6B, 0x4A))

add_heading_line("Objective")
add_para(
    "Grow as a data and AI professional — leading projects that sharpen analytics, "
    "deploy more robust models, and create experiences that matter."
)

add_heading_line("Professional Summary")
add_para(
    "Data professional from Brazil, based in Dubai, with 5+ years focused on analytics "
    "and BI — and a strong edge in practical AI solutions. Expert in Python across the "
    "full data lifecycle: extraction, transformation, visualization and automation. "
    "Designs database infrastructure, APIs and automated pipelines; builds LLM-powered "
    "tools and intelligent chatbots that turn complex data into decisions that improve "
    "performance and operational efficiency."
)

add_heading_line("Competencies")
add_para(
    "AI / LLM Solutions  |  Python  |  Data Visualization  |  Database Infrastructure  |  "
    "Automation  |  Ecommerce Specialist  |  SQL  |  ETL / Pipelines  |  APIs  |  Looker Studio  |  "
    "Chatbots  |  Frontend / Web  |  A/B Testing  |  CRM Analytics"
)

add_heading_line("Professional Experience")

add_para("Goettling Interiors — Senior Data Analyst", bold=True, space_after=0)
add_para("Dubai, UAE  |  Aug 2024 – Present", size=10, italic=True, space_after=4)
for b in [
    "Built an intelligent WhatsApp solution with AI/LLM automation, streamlining customer conversations and extracting actionable insights from chat data.",
    "Built an automated migration robot for Winner Design → Winner Flex, migrating 2,010 projects successfully and saving ~64 hours (~8 working days) of repetitive manual work.",
    "Led the frontend rebuild of the Goettling digital experience (HTML/CSS/JS), shipping interactive product journeys and performance-focused UX.",
    "Delivered the full Occhio Dubai e-commerce site end-to-end — storefront, content architecture and conversion flows — live at occhiodubai.ae.",
    "Structured business intelligence datasets in SQL and Python, turning raw operational data into actionable insights for leadership.",
    "Integrated platforms via API calls, ensuring reliable data flow across systems and business processes.",
    "Applied statistical methods, machine learning and predictive modeling to improve operational efficiency and guide strategy.",
    "Built and maintained dashboards in Looker Studio and Excel, delivering clear reports that support faster decision-making.",
    "Designed and ran A/B tests on marketing and engagement initiatives, optimizing campaigns with evidence-based results.",
]:
    add_bullet(b)

add_para("Vinci Shoes — Ecommerce Analyst", bold=True, space_after=0, space_before=8)
add_para("Porto Alegre, Brazil  |  Oct 2020 – Sep 2024", size=10, italic=True, space_after=4)
for b in [
    "Analyzed and structured BI data in Excel and analytics tools, driving performance optimization across e-commerce operations.",
    "Integrated platforms through API calls, improving data connectivity between storefront, CRM and marketing systems.",
    "Executed structured A/B tests on UX and campaigns, informing product and marketing decisions with measurable evidence.",
    "Managed CRM analytics to refine customer engagement strategies and improve campaign effectiveness.",
]:
    add_bullet(b)

add_para("Agência YouGo — Marketing Analyst", bold=True, space_after=0, space_before=8)
add_para("Porto Alegre, Brazil  |  Feb 2020 – Oct 2020", size=10, italic=True, space_after=4)
for b in [
    "Executed inbound and paid media strategies in Google Ads and social ads, optimizing campaigns for higher ROI.",
    "Tracked end-to-end campaign performance with analytics, refining targeting to deliver measurable results for clients.",
]:
    add_bullet(b)

add_para("Vnda Ecommerce — UX Designer", bold=True, space_after=0, space_before=8)
add_para("Porto Alegre, Brazil  |  May 2019 – Jan 2020", size=10, italic=True, space_after=4)
for b in [
    "Designed wireframes and high-fidelity prototypes in Figma for e-commerce sites, improving usability across industries.",
    "Collaborated with development teams to implement user-centered layouts aligned with client conversion goals.",
]:
    add_bullet(b)

add_para("Uni Explorer Turismo — Marketing Analyst", bold=True, space_after=0, space_before=8)
add_para("Porto Alegre, Brazil  |  Nov 2016 – May 2019", size=10, italic=True, space_after=4)
for b in [
    "Built buyer personas, journeys and sales funnels using inbound marketing methods, increasing engagement and conversion focus.",
    "Optimized campaigns with Facebook Ads, Google Ads, Tag Manager and Analytics, maximizing ROI through continuous measurement.",
    "Led content production teams to execute cohesive social and content strategies tied to sales outcomes.",
]:
    add_bullet(b)

add_heading_line("Languages")
add_para("English — Fluent  |  Portuguese — Native  |  Spanish — Intermediate")

add_heading_line("Education")
add_para("Postgraduate Degree, Data Science — POLI USP PRO", bold=True, space_after=0)
add_para(
    "São Paulo, Brazil  |  Apr 2024 – Apr 2026  |  Python, Big Data and applied data science.",
    size=10,
    italic=True,
    space_after=6,
)
add_para("Bachelor's Degree, Advertising — Centro Universitário Metodista IPA", bold=True, space_after=0)
add_para(
    "Porto Alegre, Brazil  |  Jan 2015 – Dec 2019  |  Data-driven marketing and data analytics focus.",
    size=10,
    italic=True,
)

doc.save(path)
print("Saved:", path)
